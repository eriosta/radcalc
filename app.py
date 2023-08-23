import streamlit as st
from docx import Document
from io import BytesIO

def generate_report(T_value, X, Y, percentage, fsm, m_values):
    # Create a new document
    doc = Document()

    # Populate the report content
    doc.add_heading('PROCEDURE', level=1)
    doc.add_paragraph(f"{st.text_input('Input Procedure:', 'GRE/SE EPI')} MR elastography and chemical shift-encoded GRE sequences were performed for liver fibrosis, fat, and iron quantification on a {T_value} Tesla scanner.")
    
    doc.add_heading('MR Liver Elastography', level=1)
    doc.add_paragraph(f"Mean liver stiffness (weighted mean of {len(m_values)} measurements): {fsm:.2f} kPa (range {min(m_values):.2f} – {max(m_values):.2f} kPa).")
    doc.add_paragraph("Interpretation of MR elastography results. Mean LSM:")
    # ... add the rest of the content similarly ...
    
    # Create an in-memory binary stream to save doc to
    stream = BytesIO()
    doc.save(stream)
    return stream

def calculate_iron(T_value, X):
    if T_value == '1.5 T':
        return 2.603E-2 * X - 0.16
    elif T_value == '2.89 T':
        return 1.400E-2 * X - 0.03
    elif T_value == '3.0 T':
        return 1.349E-2 * X - 0.03

def iron_grading(Y):
    if Y < 1.8:
        return "Grade 0: Normal. No Fe overload."
    elif 1.8 <= Y <= 3.2:
        return "Grade 1: Mild Fe overload."
    elif 3.2 <= Y <= 7:
        return "Grade 2: Moderate Fe overload."
    elif 7 <= Y <= 15:
        return "Grade 3: Severe Fe overload."
    else:
        return "Grade 4: Extreme Fe overload."

def fat_grading(percentage):
    if percentage < 5:
        return "Grade 0: Normal. No steatosis."
    elif 5 <= percentage <= 17:
        return "Grade 1: Mild steatosis."
    elif 17 <= percentage <= 22:
        return "Grade 2: Moderate steatosis."
    else:
        return "Grade 3: Severe steatosis."

def calculate_fsm(m_values, a_values):
    return sum([m*a for m,a in zip(m_values, a_values)]) / sum(a_values)

def fsm_grading(fsm):
    if fsm < 2.5:
        return "Normal or chronic inflammation"
    elif 2.5 <= fsm <= 3.0:
        return "Stage 1-2 fibrosis."
    elif 3.0 <= fsm <= 3.5:
        return "Stage 2-3 fibrosis."
    elif 3.5 <= fsm <= 4.0:
        return "Stage 3-4 fibrosis."
    else:
        return "Stage 4-5 fibrosis."

st.title('MR Elastography App')

# Calculate FSM
st.header('Calculate FSM')
m_values = [st.number_input(f'Enter ROI{i} LSM:', value=0.0) for i in range(1, 5)]
a_values = [st.number_input(f'Enter ROI{i} area:', value=0.0) for i in range(1, 5)]

if st.button('Calculate FSM'):
    if sum(a_values) == 0:
        st.error("Sum of ROI areas cannot be zero!")
    else:
        fsm = calculate_fsm(m_values, a_values)
        st.write(f"FSM: {fsm} kPa")
        st.write(fsm_grading(fsm))

# Determine Steatosis Grade
st.header('Determine Steatosis Grade')
percentage = st.slider('Enter PDFF:', min_value=0.0, max_value=100.0, value=0.0)
if st.button('Calculate Steatosis Grade'):
    st.write(fat_grading(percentage))

# Calculate LIC
st.header('Calculate LIC')
T_values = ['1.5 T', '2.89 T', '3.0 T']
T_value = st.selectbox('Select MR:', T_values)
X = st.number_input('Enter R2* Value:', value=0.0)

if st.button('Calculate LIC'):
    Y = calculate_iron(T_value, X)
    st.write(f"LIC: {Y} mg/g")
    st.write(iron_grading(Y))

# Clear results
st.header('Clear Results')
if st.button('Clear'):
    st.experimental_rerun()

# Section to generate report
st.header('Generate Report')
if st.button('Generate Report'):
    stream = generate_report(T_value, X, Y, percentage, fsm, m_values)
    buf = BytesIO(stream.getvalue())
    st.download_button(
        label="Download Report",
        data=buf,
        file_name='MRElastographyReport.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )